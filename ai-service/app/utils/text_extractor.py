import io
import pdfplumber
import docx
from fastapi import UploadFile


async def extract_text_from_file(file: UploadFile) -> str:
    content = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text += page_text + "\n"
        return "\n".join(line.strip() for line in text.split("\n") if line.strip())

    if filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)

    raise ValueError("Unsupported file format. Upload a PDF or DOCX file.")
